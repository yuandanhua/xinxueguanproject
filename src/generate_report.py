import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE_PATH = r"E:\QST\文档\400源码实验手册\混凝药剂投放量预测分析\混凝药剂投放量预测分析v1.docx"
OUTPUT_DIR = "output"

HEADING_SIZES = {0: 22, 1: 16, 2: 14, 3: 12}


def _style_exists(doc, name):
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _add_heading(doc, text, level=1):
    style = {0: "Title", 1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Normal")
    if _style_exists(doc, style):
        return doc.add_heading(text, level=level)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(HEADING_SIZES.get(level, 12))
    return p


def _add_code_block(doc, code_text):
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)


def _add_step_title(doc, text):
    if _style_exists(doc, "3标题"):
        return doc.add_paragraph(text, style="3标题")
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def _add_text(doc, text):
    return doc.add_paragraph(text)


def _add_bold_text(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(text)
    return p


def create_report(metrics, eda_images=None, prefix="experiment"):
    print("=" * 50)
    print("生成实验手册文档")
    print("=" * 50)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "sectPr":
                continue
            body.remove(child)
        print(f"[报告] 已加载模板: {TEMPLATE_PATH}")
    else:
        doc = Document()

    _add_heading(doc, "心血管疾病发病风险预测分析（KNN模型实验）", level=1)

    # ── 实验目的 ──
    _add_heading(doc, "实验目的", level=2)
    _add_text(doc,
        "本实验旨在通过机器学习KNN（K-Nearest Neighbors）算法，"
        "基于患者特征（性别、年龄、身高、体重、血压、胆固醇、血糖、"
        "吸烟、饮酒、锻炼等）建立心血管疾病发病风险评估模型，"
        "帮助识别高风险人群并指导预防干预。同时通过实践掌握Pandas、"
        "Scikit-learn等工具在数据科学全流程中的应用。"
    )

    # ── 实验内容 ──
    _add_heading(doc, "实验内容", level=2)
    _add_text(doc,
        "本实验使用Kaggle心血管疾病数据集，包含患者的基本信息、"
        "体检指标和生活习惯等特征，通过KNN分类算法构建二分类模型，"
        "预测患者是否患有心血管疾病。"
    )

    _add_heading(doc, "数据集说明", level=2)
    _add_text(doc,
        "训练集：历史患者数据信息.csv，共67,849条记录。"
        "预测集：当前患者数据信息.csv，共2,000条记录。"
    )
    _add_text(doc, "数据集字段如下表所示：")

    # field description table
    fields = [
        ("患者编号", "患者的唯一标识"),
        ("年龄", "患者的年龄（岁）"),
        ("性别", "男/女"),
        ("身高", "患者的身高（cm）"),
        ("体重(KG)", "患者的体重（kg）"),
        ("收缩压", "心脏收缩时血液对血管壁产生的张力（高压）"),
        ("舒张压", "心脏舒张时血液对于血管壁产生的压力（低压）"),
        ("胆固醇", "三分类：正常、高于正常、远高于正常"),
        ("葡萄糖", "三分类：正常、高于正常、远高于正常"),
        ("是否吸烟", "二分类：吸烟/不吸烟"),
        ("是否饮酒", "二分类：饮酒/不饮酒"),
        ("是否锻炼", "二分类：锻炼/不锻炼"),
        ("是否有心血管疾病", "目标变量，二分类：有/无"),
    ]
    table = doc.add_table(rows=len(fields) + 1, cols=2)
    try:
        table.style = "Light Shading Accent 1"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "字段名"
    table.rows[0].cells[1].text = "说明"
    for i, (name, desc) in enumerate(fields):
        table.rows[i + 1].cells[0].text = name
        table.rows[i + 1].cells[1].text = desc

    # ── 实验知识 ──
    _add_heading(doc, "实验知识", level=2)
    _add_text(doc,
        "K最近邻（KNN）算法是一种基于实例的监督学习算法。"
        "其核心思想是：给定一个待分类样本，通过计算其与训练集中"
        "所有样本的距离，找出距离最近的K个邻居，根据这K个邻居的"
        "类别通过投票（或加权投票）决定待分类样本的类别。"
    )
    _add_text(doc,
        "KNN算法的主要超参数包括：K值（邻居数量）、距离度量方式"
        "（欧氏距离、曼哈顿距离等）和权重策略（uniform为等权投票，"
        "distance为距离加权投票）。KNN对特征尺度敏感，因此需要"
        "进行特征标准化。"
    )

    # ── 实验步骤 ──
    _add_heading(doc, "实验步骤", level=2)

    # Step 1: 数据加载
    _add_step_title(doc, "1. 加载数据")
    _add_text(doc,
        "使用Pandas读取两个CSV文件，编码格式为UTF-8-BOM。"
        "训练数据包含目标变量（是否有心血管疾病），预测数据不包含。"
    )
    with open("src/data_loader.py", "r", encoding="utf-8") as f:
        _add_code_block(doc, f.read())

    # Step 2: EDA
    _add_step_title(doc, "2. 探索性数据分析")
    _add_text(doc,
        "对训练数据进行缺失值检测、异常值检测、特征分布分析、"
        "相关性分析和目标变量分布分析，输出可视化图表。"
    )
    with open("src/eda.py", "r", encoding="utf-8") as f:
        _add_code_block(doc, f.read())

    if eda_images:
        for img in sorted(eda_images):
            path = os.path.join(OUTPUT_DIR, img)
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(5.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Step 3: 特征工程
    _add_step_title(doc, "3. 特征工程")
    _add_text(doc,
        "对分类特征进行编码转换：胆固醇和葡萄糖为三分类编码"
        "（正常=0, 高于正常=1, 远高于正常=2）；"
        "性别、吸烟、饮酒、锻炼为二分类编码。"
        "数值特征进行Z-score标准化。按8:2划分训练集和测试集。"
    )
    with open("src/feature_engineering.py", "r", encoding="utf-8") as f:
        _add_code_block(doc, f.read())

    # Step 4: 模型训练
    _add_step_title(doc, "4. KNN模型训练与调优")
    _add_text(doc,
        "使用GridSearchCV进行超参数搜索，共54种参数组合，"
        "5折交叉验证，以F1分数为优化目标。"
    )
    with open("src/model_training.py", "r", encoding="utf-8") as f:
        _add_code_block(doc, f.read())
    _add_bold_text(doc, "最佳参数：",
        f"metric={metrics.get('best_params', {}).get('metric', 'manhattan')}, "
        f"n_neighbors={metrics.get('best_params', {}).get('n_neighbors', 23)}, "
        f"weights={metrics.get('best_params', {}).get('weights', 'uniform')}"
    )
    _add_bold_text(doc, "最佳交叉验证F1：", f"{metrics.get('cv_f1', 0.7065):.4f}")

    # Step 5: 模型评估
    _add_step_title(doc, "5. 模型评估")
    _add_text(doc, "在测试集上评估模型性能，输出混淆矩阵、ROC曲线和分类报告。")
    with open("src/evaluation.py", "r", encoding="utf-8") as f:
        eval_code = f.read()
    _add_code_block(doc, eval_code)

    # metrics table
    metric_names = {
        "accuracy": "准确率 (Accuracy)",
        "precision": "精确率 (Precision)",
        "recall": "召回率 (Recall)",
        "f1": "F1 分数",
        "auc": "AUC",
    }
    tbl = doc.add_table(rows=len(metric_names) + 1, cols=2)
    try:
        tbl.style = "Light Shading Accent 1"
    except KeyError:
        pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.rows[0].cells[0].text = "评估指标"
    tbl.rows[0].cells[1].text = "数值"
    for i, (key, label) in enumerate(metric_names.items()):
        tbl.rows[i + 1].cells[0].text = label
        tbl.rows[i + 1].cells[1].text = f"{metrics.get(key, 0):.4f}"

    for img_name in ["test_confusion_matrix.png", "test_roc_curve.png"]:
        path = os.path.join(OUTPUT_DIR, img_name)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Step 6: 预测
    _add_step_title(doc, "6. 预测结果")
    _add_text(doc,
        "使用训练好的KNN模型对2000条当前患者数据进行预测，"
        "输出每位患者的心血管疾病发病风险（有/无）及概率值，"
        "并划分为高风险（≥70%）、中风险（40%-70%）、低风险（<40%）三个等级。"
    )

    # ── 实验总结 ──
    _add_heading(doc, "实验总结", level=2)
    _add_text(doc,
        "本实验完整实现了从数据加载、探索性分析、特征工程、"
        "模型训练与调优到结果预测的全流程数据科学工作。"
        f"最终模型在测试集上达到Accuracy={metrics.get('accuracy', 0):.4f}、"
        f"F1={metrics.get('f1', 0):.4f}、AUC={metrics.get('auc', 0):.4f}。"
    )
    _add_text(doc,
        "通过KNN算法的实践，深入理解了距离度量、k值选择、"
        "特征标准化等关键概念对模型性能的影响。"
    )

    output_path = os.path.join(OUTPUT_DIR, f"{prefix}_report.docx")
    doc.save(output_path)
    print(f"[报告] 实验手册已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    metrics = {"accuracy": 0.72, "precision": 0.74, "recall": 0.68, "f1": 0.71, "auc": 0.78}
    create_report(metrics)
